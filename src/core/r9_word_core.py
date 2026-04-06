# src/core/r9_word_core.py
# ============================================================
# R9 — Reporte Word mensual profesional
# Sis Rodríguez Contadores Públicos
# ============================================================

from __future__ import annotations
from pathlib import Path
from datetime import datetime
import json
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paleta corporativa ────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1F, 0x38, 0x64)
AZUL_MEDIO   = RGBColor(0x2E, 0x75, 0xB6)
AZUL_CLARO   = RGBColor(0xD6, 0xE4, 0xF0)
GRIS_TEXTO   = RGBColor(0x40, 0x40, 0x40)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
VERDE_OK     = RGBColor(0x1E, 0x8B, 0x4C)
NARANJA_PEND = RGBColor(0xD6, 0x7D, 0x00)

LOGO_PATH    = Path("src/assets/logo_sisrodriguez.png")

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}

# ── Helpers de formato ────────────────────────────────────────────────────────
def _fmt_mxn(valor) -> str:
    try:
        return f"$ {float(valor):>14,.2f}"
    except Exception:
        return "$ 0.00"

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _set_cell_borders(cell, color="2E75B6", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    size)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _bold_run(para, text: str, size=10, color=None, italic=False):
    run       = para.add_run(text)
    run.bold  = True
    run.font.size  = Pt(size)
    run.font.name  = "Calibri"
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run

def _normal_run(para, text: str, size=10, color=None, bold=False):
    run      = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    return run

def _add_divider(doc, color="2E75B6"):
    p  = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p

def _set_table_style(table):
    """Aplica estilo limpio a tabla — sin estilo heredado."""
    table.style = "Table Grid"

# ── Carga de datos ────────────────────────────────────────────────────────────
def _load_sheet(path: Path, sheet="CFDI") -> pd.DataFrame:
    try:
        return pd.read_excel(path, sheet_name=sheet)
    except Exception:
        return pd.DataFrame()

def _stats(df: pd.DataFrame) -> dict:
    if df.empty:
        return dict(total=0.0, subtotal=0.0, iva=0.0,
                    total_pue=0.0, count=0, count_pue=0,
                    count_ppd=0, total_ppd=0.0,
                    por_tipo={})
    mp  = df["METODO_PAGO"].astype(str).str.upper()
    pue = df[mp.eq("PUE")]
    ppd = df[mp.eq("PPD")]

    por_tipo = {}
    for t in df["TIPO_COMPROB"].unique():
        sub = df[df["TIPO_COMPROB"] == t]
        por_tipo[t] = {
            "count":    len(sub),
            "subtotal": float(round(sub["SUBTOTAL_MXN"].sum(), 2)) if "SUBTOTAL_MXN" in sub else 0.0,
            "iva":      float(round(sub["IVA_16_CALC"].sum(), 2))  if "IVA_16_CALC"  in sub else 0.0,
            "total":    float(round(sub["TOTAL"].sum(), 2))         if "TOTAL"         in sub else 0.0,
        }

    return dict(
        count     = len(df),
        subtotal  = float(round(df["SUBTOTAL_MXN"].sum(), 2)) if "SUBTOTAL_MXN" in df else 0.0,
        iva       = float(round(df["IVA_16_CALC"].sum(), 2))  if "IVA_16_CALC"  in df else 0.0,
        total     = float(round(df["TOTAL"].sum(), 2))         if "TOTAL"         in df else 0.0,
        count_pue = len(pue),
        total_pue = float(round(pue["TOTAL"].sum(), 2)) if not pue.empty and "TOTAL" in pue else 0.0,
        count_ppd = len(ppd),
        total_ppd = float(round(ppd["TOTAL"].sum(), 2)) if not ppd.empty and "TOTAL" in ppd else 0.0,
        por_tipo  = por_tipo,
    )

def _load_pagos(path: Path) -> pd.DataFrame:
    try:
        return pd.read_excel(path, sheet_name="PAGOS")
    except Exception:
        return pd.DataFrame()


# ── Bloque: encabezado con logo ───────────────────────────────────────────────
def _add_header(doc: Document, rfc: str, nombre: str, yyyy_mm: str):
    year, month = yyyy_mm.split("-")
    mes_label   = MESES_ES.get(int(month), month).capitalize()
    periodo     = f"{mes_label} {year}"

    # Tabla de encabezado: logo | datos
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(12)
    row = tbl.rows[0]

    # Celda logo
    cell_logo = row.cells[0]
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para_logo = cell_logo.paragraphs[0]
    para_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        run = para_logo.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(4.2))
    else:
        _bold_run(para_logo, "SIS RODRÍGUEZ\nContadores Públicos",
                  size=11, color=AZUL_OSCURO)

    # Celda datos del reporte
    cell_info = row.cells[1]
    cell_info.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell_info, "1F3864")

    p1 = cell_info.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(2)
    p1.paragraph_format.left_indent  = Cm(0.3)
    _bold_run(p1, "REPORTE MENSUAL DE CFDI", size=13, color=BLANCO)

    p2 = cell_info.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.left_indent = Cm(0.3)
    p2.paragraph_format.space_after = Pt(2)
    _bold_run(p2, f"{nombre}", size=11, color=AZUL_CLARO)

    p3 = cell_info.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.left_indent = Cm(0.3)
    p3.paragraph_format.space_after = Pt(4)
    _normal_run(p3, f"RFC: {rfc}   |   Período: {periodo}", size=10, color=BLANCO)

    doc.add_paragraph()  # espacio


# ── Bloque: sección con título ────────────────────────────────────────────────
def _section_title(doc: Document, texto: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    _bold_run(p, f"  {texto}", size=11, color=BLANCO)
    # Fondo azul en el párrafo
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "2E75B6")
    pPr.append(shd)


# ── Bloque: tabla resumen fiscal ──────────────────────────────────────────────
def _add_tabla_resumen_fiscal(doc: Document, stats: dict, rol: str):
    TIPOS_LABEL = {
        "I": "Ingresos (Facturas)",
        "N": "Nómina",
        "T": "Traslados",
        "E": "Egresos / Notas de crédito",
        "P": "Complementos de Pago",
    }
    por_tipo = stats.get("por_tipo", {})
    tipos_presentes = [t for t in ["I", "N", "T", "E", "P"] if t in por_tipo]

    if not tipos_presentes:
        p = doc.add_paragraph()
        _normal_run(p, "  Sin movimientos en el período.", size=10, color=GRIS_TEXTO)
        return

    # Encabezados de tabla
    headers = ["Concepto", "No. CFDI", "Subtotal", "IVA (16%)", "Total"]
    col_w   = [Cm(5.5), Cm(2.0), Cm(3.5), Cm(3.0), Cm(3.5)]

    tbl = doc.add_table(rows=1 + len(tipos_presentes) + 1, cols=5)
    _set_table_style(tbl)

    # Fila de encabezado
    hdr_row = tbl.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_w)):
        cell = hdr_row.cells[i]
        cell.width = w
        _set_cell_bg(cell, "2E75B6")
        _set_cell_borders(cell, "FFFFFF", "2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(p, h, size=9, color=BLANCO)

    # Filas de datos por tipo
    totales = {"No. CFDI": 0, "Subtotal": 0.0, "IVA (16%)": 0.0, "Total": 0.0}
    for r_idx, tipo in enumerate(tipos_presentes, start=1):
        d     = por_tipo[tipo]
        label = TIPOS_LABEL.get(tipo, tipo)
        vals  = [label, d["count"], d["subtotal"], d["iva"], d["total"]]
        bg    = "D6E4F0" if r_idx % 2 == 0 else "FFFFFF"
        row_  = tbl.rows[r_idx]
        for c_idx, val in enumerate(vals):
            cell = row_.cells[c_idx]
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell, "BDC3C7", "2")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx >= 1 else WD_ALIGN_PARAGRAPH.LEFT
            text = _fmt_mxn(val) if c_idx >= 2 else str(val)
            _normal_run(p, text, size=9)
        totales["No. CFDI"]  += d["count"]
        totales["Subtotal"]  += d["subtotal"]
        totales["IVA (16%)"] += d["iva"]
        totales["Total"]     += d["total"]

    # Fila total general
    total_row_ = tbl.rows[-1]
    total_vals = ["TOTAL GENERAL", totales["No. CFDI"],
                  totales["Subtotal"], totales["IVA (16%)"], totales["Total"]]
    for c_idx, val in enumerate(total_vals):
        cell = total_row_.cells[c_idx]
        _set_cell_bg(cell, "1F3864")
        _set_cell_borders(cell, "FFFFFF", "2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx >= 1 else WD_ALIGN_PARAGRAPH.LEFT
        text = _fmt_mxn(val) if c_idx >= 2 else str(val)
        _bold_run(p, text, size=9, color=BLANCO)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Bloque: tabla de pagos pendientes ─────────────────────────────────────────
def _add_tabla_pagos(doc: Document, df_pagos: pd.DataFrame):
    if df_pagos.empty:
        p = doc.add_paragraph()
        _normal_run(p, "  No se encontraron complementos de pago en el período.", size=10)
        return

    headers  = ["Fecha Pago", "Cliente / Proveedor", "Forma de Pago",
                 "Importe Pagado", "Saldo Insoluto", "Estado"]
    col_w    = [Cm(2.5), Cm(5.0), Cm(2.8), Cm(2.8), Cm(2.5), Cm(2.0)]

    n_rows   = min(len(df_pagos), 20)  # máximo 20 en el Word, el Excel tiene todos
    tbl      = doc.add_table(rows=1 + n_rows, cols=6)
    _set_table_style(tbl)

    # Encabezado
    for i, (h, w) in enumerate(zip(headers, col_w)):
        cell = tbl.rows[0].cells[i]
        cell.width = w
        _set_cell_bg(cell, "1A5276")
        _set_cell_borders(cell, "FFFFFF", "2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bold_run(p, h, size=8, color=BLANCO)

    for r_idx, (_, pago) in enumerate(df_pagos.head(n_rows).iterrows(), start=1):
        saldo_ins = float(pago.get("SALDO_INSOLUTO", 0) or 0)
        pagado    = saldo_ins == 0
        bg        = "EAFAF1" if pagado else "FEF9E7"
        estado    = "Pagado" if pagado else "Pendiente"
        color_est = VERDE_OK if pagado else NARANJA_PEND

        vals = [
            str(pago.get("FECHA_PAGO", ""))[:10],
            str(pago.get("NOMBRE_CLIENTE", ""))[:40],
            str(pago.get("FORMA_PAGO_DESC", "")),
            _fmt_mxn(pago.get("IMPORTE_PAGADO", 0)),
            _fmt_mxn(saldo_ins),
            estado,
        ]
        row_ = tbl.rows[r_idx]
        for c_idx, val in enumerate(vals):
            cell = row_.cells[c_idx]
            _set_cell_bg(cell, bg)
            _set_cell_borders(cell, "BDC3C7", "2")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx in (3, 4) else WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 5) else WD_ALIGN_PARAGRAPH.LEFT
            if c_idx == 5:
                _bold_run(p, val, size=8, color=color_est)
            else:
                _normal_run(p, val, size=8)

    if len(df_pagos) > n_rows:
        p = doc.add_paragraph()
        _normal_run(p, f"  * Se muestran los primeros {n_rows} registros. Ver hoja PAGOS en Excel para el detalle completo.",
                    size=8, color=GRIS_TEXTO, bold=False)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Bloque: pie de página ─────────────────────────────────────────────────────
def _add_footer_note(doc: Document, yyyy_mm: str):
    _add_divider(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    _normal_run(p,
        f"Sis Rodríguez Contadores Públicos  •  Reporte generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}  •  Uso interno",
        size=8, color=GRIS_TEXTO)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _normal_run(p2,
        "Este documento es confidencial y está destinado exclusivamente al uso del despacho contable.",
        size=8, color=GRIS_TEXTO)


# ── Builder principal ─────────────────────────────────────────────────────────
def build_month_summary_docx(
    exports_dir: Path,
    rfc: str,
    yyyy_mm: str,
    nombre_empresa: str = "",
) -> Path:
    base = exports_dir / rfc / yyyy_mm
    rec  = base / f"{rfc}_{yyyy_mm}_RECIBIDAS_Facturas.xlsx"
    emi  = base / f"{rfc}_{yyyy_mm}_EMITIDAS_Facturas.xlsx"

    df_r = _load_sheet(rec)
    df_e = _load_sheet(emi)
    s_rec = _stats(df_r)
    s_emi = _stats(df_e)

    # Pagos de ambos archivos
    df_pagos_r = _load_pagos(rec)
    df_pagos_e = _load_pagos(emi)
    df_pagos   = pd.concat([df_pagos_r, df_pagos_e], ignore_index=True)
    if not df_pagos.empty and "FECHA_PAGO" in df_pagos.columns:
        df_pagos = df_pagos.sort_values("FECHA_PAGO")

    # Nombre de la empresa desde el Excel si no se pasa
    if not nombre_empresa and not df_e.empty and "EMISOR_NOMBRE" in df_e.columns:
        nombre_empresa = df_e["EMISOR_NOMBRE"].dropna().iloc[0] if len(df_e) > 0 else rfc
    if not nombre_empresa and not df_r.empty and "RECEPTOR_NOMBRE" in df_r.columns:
        nombre_empresa = df_r["RECEPTOR_NOMBRE"].dropna().iloc[0] if len(df_r) > 0 else rfc
    if not nombre_empresa:
        nombre_empresa = rfc

    # ── Crear documento ───────────────────────────────────────────────────────
    doc = Document()

    # Márgenes de página
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Fuente por defecto
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Encabezado con logo ───────────────────────────────────────────────────
    _add_header(doc, rfc, nombre_empresa, yyyy_mm)

    # ── FACTURAS EMITIDAS ─────────────────────────────────────────────────────
    _section_title(doc, f"FACTURAS EMITIDAS  ({s_emi['count']} CFDIs)")
    _add_tabla_resumen_fiscal(doc, s_emi, "EMITIDAS")

    # Datos rápidos PUE/PPD
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    _normal_run(p, f"PUE (pago en una exhibición): ", size=9, bold=True)
    _normal_run(p, f"{s_emi['count_pue']} CFDIs  |  {_fmt_mxn(s_emi['total_pue'])}", size=9)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(6)
    _normal_run(p2, f"PPD (pago en parcialidades): ", size=9, bold=True)
    _normal_run(p2, f"{s_emi['count_ppd']} CFDIs  |  {_fmt_mxn(s_emi['total_ppd'])}", size=9)

    # ── FACTURAS RECIBIDAS ────────────────────────────────────────────────────
    _section_title(doc, f"FACTURAS RECIBIDAS  ({s_rec['count']} CFDIs)")
    _add_tabla_resumen_fiscal(doc, s_rec, "RECIBIDAS")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    _normal_run(p, f"PUE: ", size=9, bold=True)
    _normal_run(p, f"{s_rec['count_pue']} CFDIs  |  {_fmt_mxn(s_rec['total_pue'])}", size=9)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(6)
    _normal_run(p2, f"PPD: ", size=9, bold=True)
    _normal_run(p2, f"{s_rec['count_ppd']} CFDIs  |  {_fmt_mxn(s_rec['total_ppd'])}", size=9)

    # ── COMPLEMENTOS DE PAGO ─────────────────────────────────────────────────
    _section_title(doc, f"COMPLEMENTOS DE PAGO  ({len(df_pagos)} registros)")
    _add_tabla_pagos(doc, df_pagos)

    # ── BALANCE RÁPIDO ────────────────────────────────────────────────────────
    _section_title(doc, "BALANCE DEL PERÍODO")
    tbl_b = doc.add_table(rows=4, cols=2)
    _set_table_style(tbl_b)
    balance_data = [
        ("Total Emitido (ingresos)",  s_emi["total"]),
        ("Total Recibido (gastos)",   s_rec["total"]),
        ("IVA Emitido",               s_emi["iva"]),
        ("IVA Recibido (acreditable)", s_rec["iva"]),
    ]
    for r_idx, (label, val) in enumerate(balance_data):
        row_ = tbl_b.rows[r_idx]
        bg   = "D6E4F0" if r_idx % 2 == 0 else "FFFFFF"
        _set_cell_bg(row_.cells[0], bg)
        _set_cell_bg(row_.cells[1], bg)
        _set_cell_borders(row_.cells[0], "BDC3C7", "2")
        _set_cell_borders(row_.cells[1], "BDC3C7", "2")
        row_.cells[0].width = Cm(9)
        row_.cells[1].width = Cm(4)
        p0 = row_.cells[0].paragraphs[0]
        p1 = row_.cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _bold_run(p0, label, size=9)
        _bold_run(p1, _fmt_mxn(val), size=9, color=AZUL_OSCURO)

    doc.add_paragraph()

    # ── Pie de página ─────────────────────────────────────────────────────────
    _add_footer_note(doc, yyyy_mm)

    # ── Guardar ───────────────────────────────────────────────────────────────
    base.mkdir(parents=True, exist_ok=True)
    out = base / f"RESUMEN_{rfc}_{yyyy_mm}.docx"
    doc.save(str(out))

    # Manifest actualizado
    (base / "manifest.json").write_text(
        json.dumps({
            "rfc": rfc, "yyyy_mm": yyyy_mm, "nombre": nombre_empresa,
            "generado": datetime.now().isoformat(),
            "files": {
                "RECIBIDAS": rec.name if rec.exists() else None,
                "EMITIDAS":  emi.name if emi.exists() else None,
                "RESUMEN":   out.name,
            }
        }, ensure_ascii=False, indent=2),
        encoding="utf-8"
    )

    return out